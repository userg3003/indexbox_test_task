from pathlib import Path

import sqlite3
import pandas as pd
import numpy
import docx


def cagr(start: pd.DataFrame, end: pd.DataFrame, period: int = 2):
    """
    CAGR (англ. Compound annual growth rate) — совокупный среднегодовой темп роста. Выражается в процентах и показывает,
     на сколько процентов за год прирастает изучаемый параметр.
    :param start: - значение на начало периода
    :param end: - значение на конец периода
    :param period: - период
    :return:
    """
    return ((end / start) ** (1 / period) - 1) * 100


def make_rows_bold(*rows):
    """
    Сделать текст в ячейке жирным
    :param rows: "массив" ячеек
    :return: None
    """
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True


table = "testidprod"

db_path = Path.cwd()
db_name = "test.db"
full_path = db_path.joinpath(db_name)

connection = sqlite3.connect(full_path)
cursor = connection.cursor()

"""
Предварительные замечания:
Получить данные из БД и заполнить дата-фрэйм можно несколькими способами:

- Считать все данные в переменную rows: 
rows = cursor.execute(f"
select factor, country, year, res from main.{table} where
    factor in (1,2) and
    bs = 0 and
    partner is null  and
    state is null 
    ").fetchall()
    
и выполнить дальнейшие расчёты с применением непосредственно Python, либо используя 
библиотеку Pandas.
 
- Считать данные непосредственно в  дата-фрэйм, при этом выполнив предварительные 
вычисления сумм средствами БД.

query = f"select factor  as Factor, year as Year, sum(res) as world from {table} group by factor, year" 
     "having partner is null and state is null and bs == 0 and factor in (1, 2);"

df_src = pd.read_sql_query(query, connection)


Выбор конкретного варианта следует делать исходя из конкретной ситуации (ограничения по памяти, 
расположение и тип БД). 
В данном случае в качестве БД выбрана SQLite3 и ограничения по памяти не установлены.
Хотя 1-й вариант (без вычисления сумм) выполняется незначительно быстрее 2-го, расходы на дальнейшие 
вычисления сводят его преимущества к 0.

В связи с этим выбран 2-й вариант 

"""

query = f"""
     select factor  as Factor, year as Year, sum(res) as world from {table} group by factor, year 
     having partner is null and state is null and bs == 0 and factor in (1, 2);
 """

df_src = pd.read_sql_query(query, connection)

cursor.close()
"""
Предварительный анализ данных в БД показал, что отсутствуют данные за 2006 и 2020 гг.
По заданию эти данные должны присутствовать (явно не указано, но в таблице с примером 
результата эти года присутствуют) 
"""

#  Добавить отсутствующие года
empty_years = {"Factor": [1, 1, 2, 2],
               "Year": [2006, 2020, 2006, 2020],
               "world": [numpy.NaN, numpy.NaN, numpy.NaN, numpy.NaN]}
df_empty_years = pd.DataFrame(empty_years)
df_dst = pd.concat([df_empty_years, df_src], ignore_index=True)
df_dst.sort_values(by=['Factor', 'Year'], inplace=True)
df_dst.index = pd.Series(range(len(df_dst)))
df_working = df_dst.T

# Добавить Фактор 6
for index in range(15):
    df_working.insert(30 + index, 30 + index,
                      [6, df_working[index]['Year'], df_working[index + 15]["world"] / df_working[index]["world"]])

df_working.to_excel("report.xlsx")

df_factor6 = df_working.iloc[2, list(range(30, 45))]
df_factor6 = df_factor6.reset_index()
del df_factor6['index']
df_for_report = pd.DataFrame(df_factor6)

# Выполнить вычисление по фактору 6 его рост/снижение (в среднем) за каждый год
df_for_report['cagr'] = cagr(df_for_report - df_for_report.diff().fillna(df_for_report), df_for_report, period=1)
df_for_report.loc[1, 'cagr'] = numpy.nan
df_for_report.insert(0, "Year", range(2006, 2021))

# Заполнить таблицу для отчёта
rows = df_for_report.shape[0]
cols = df_for_report.shape[1]
factor = 6

doc = docx.Document()
table = doc.add_table(rows + 1, cols + 1)
table.cell(0, 0).text = "Factor"
table.cell(0, 1).text = "Year"
table.cell(0, 2).text = "World Value"

make_rows_bold(table.rows[0])

for i in range(rows):
    table.cell(i + 1, 0).text = str(factor)
    table.cell(i + 1, 1).text = str(df_for_report.loc[i, "Year"])
    table.cell(i + 1, 2).text = str(round(df_for_report.values[i, 1], 2))

merged_cell_start = table.cell(1, 0)
merged_cell_end = table.cell(rows, 0)
cell_factor6 = merged_cell_start.merge(merged_cell_end)
cell_factor6.text = str(factor)
cell_factor6.vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER

# Вычисление CAGR по всему диапозону выполняется с учётом отсутствия данных по 2006 и 2020 годам.
condition = round(cagr(df_for_report.values[1, 1], df_for_report.values[rows - 2, 1], 12), 2)

cagr6 = "grew" if condition > 0 else "decreased"
doc.add_paragraph(
    f'Factor 6 {cagr6} by avg {condition}% every year from {int(df_for_report.values[1, 0])} to {int(df_for_report.values[rows - 2, 0])}')
doc.save('report.docx')
